#!/usr/bin/env
# -*- coding: utf-8 -*-
# poweko.com
# pip install python-docx
# PayPal!F84

from docx import Document

#document = Document('write.docx')

import re
from docx import Document


def docx_replace_regex(doc_obj, regex, replace):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex, replace)


regex1 = re.compile(r"powenko")
replace1 = r"python"
filename = "write.docx"
doc = Document(filename)
docx_replace_regex(doc, regex1, replace1)
doc.save('result1.docx')

